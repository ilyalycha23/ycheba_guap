#!/usr/bin/env python3
"""Build LR9 movies lab report (docx) from GUAP template."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
TEMPLATE = next(Path(r"d:/ycheba_guap").rglob("4321_*.docx"))
OUT = ROOT / "Отчет_ЛР9_Фильмы.docx"

# Actual stats from data/данные.txt after Unit_0 run
N_TOTAL = 3354
N_TRAIN = 2180
N_TEST = 671
N_VALID = 503
N_TEST_VALID = 1174
N_SCORE = 459
N_HITS_SCORE_PRED = 341
PROB1_MEAN = 0.7310


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def replace_in_doc(doc: Document) -> None:
    """Targeted replacements in paragraphs not fully rewritten."""
    reps = [
        ("spotify.mdb", "movies.mdb"),
        ("Spotify dataset", "Movies dataset"),
        ("prepared_spotify_score.csv", "prepared_movies_score.csv"),
        ("3. Консолидация: ref_artist_regions.csv", "3. Консолидация: регионы режиссёров"),
        ("4. Join artist → region", "4. Join режиссёр → region"),
        ("↶Spotify dataset (публичный)", "↶Movies dataset (публичный)"),
        ("1. Импорт из Access (spotify_songs)", "1. Импорт из Access (фильмы 2025)"),
        ("9. Отбор popularity ≥ 30", "9. Отбор TMDB popularity ≥ 30"),
        ("8. IsHit (streams ≥ порога)", "8. IsHit (кассовые сборы ≥ порога)"),
    ]
    for p in doc.paragraphs:
        t = p.text
        for old, new in reps:
            t = t.replace(old, new)
        if t != p.text:
            set_paragraph_text(p, t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                for old, new in reps:
                    t = t.replace(old, new)
                if t != cell.text:
                    cell.text = t


def rewrite_body(doc: Document) -> None:
    """Replace main narrative blocks with movies-themed text."""
    blocks: dict[int, str] = {
        9: (
            "Пройти полный цикл аналитического проекта в среде Loginom на данных о "
            "фильмах 2025 года: подключиться к собственной базе данных Microsoft Access, "
            "получить выборку SQL-запросом, обогатить её справочником регионов режиссёров, "
            "оценить и улучшить качество данных, подготовить признаки к моделированию, "
            "обучить и протестировать модели машинного обучения, выполнить скоринг на новых "
            "записях, оценить качество предсказаний и оформить результаты в виде наглядных "
            "отчётов. В качестве целевого признака использована бинарная метка IsHit — "
            "признак блокбастера, когда мировые кассовые сборы (Worldwide Gross) превышают "
            "порог ста миллионов. Основная модель — Local Outlier Factor (LOF) в режиме "
            "novelty; дополнительно реализована логистическая регрессия sklearn для сравнения "
            "подходов и скоринга отдельного файла с новыми фильмами."
        ),
        20: (
            "Сначала были подготовлены данные в модуле «Подготовка выборок», "
            "представленном на рисунке 1."
        ),
        25: (
            "Для хранения основной витрины использована локальная база MS Access "
            "data/movies.mdb с таблицей записей о фильмах 2025 года. В сценарии подготовки "
            "данных настроено постоянное подключение к этой базе и ссылка на него для "
            "повторного использования во всех последующих узлах импорта. Такой вариант "
            "соответствует требованию работы с собственной базой данных средствами Loginom."
        ),
        26: (
            "Из таблицы фильмов в базе movies.mdb выполняется SQL-запрос на выборку всех "
            "полей. На выходе получается набор с полями режиссёра, названия фильма, "
            "кассовых сборов, популярности TMDB, активности в соцсетях и служебными "
            "идентификаторами. Этот поток далее проходит обогащение и очистку."
        ),
        30: (
            "Кратко по блокам на рисунке 1. Узел «MS Access» задаёт параметры подключения "
            "к файлу базы. Узел «↶MS Access» — ссылка на то же подключение. Узел "
            "«1. Импорт из Access (фильмы 2025)» выполняет SQL-запрос и отдаёт таблицу "
            "в сценарий."
        ),
        33: (
            "После импорта основная таблица объединяется со справочником регионов режиссёров "
            "из CSV-файла ref_artist_regions.csv. Справочник содержит поля artist, region "
            "и market_size. Соединение выполнено по полю режиссёра: в основной таблице — "
            "Artist, в справочнике — artist. Тип соединения — LEFT JOIN: все фильмы из "
            "витрины сохраняются, регион и размер рынка подставляются там, где режиссёр "
            "найден в справочнике."
        ),
        37: (
            "Кратко по блокам на рисунке 2. «3. Консолидация: регионы режиссёров» — "
            "импорт справочника. «4. Join режиссёр → region» — объединение главной "
            "таблицы и справочника по ключу Artist / artist."
        ),
        40: (
            "На этапе до агрессивной обработки сформирован контрольный снимок "
            "qc_before_cleaning.txt — выгрузка таблицы для визуальной проверки «сырых» "
            "значений после join. В данных видны характерные проблемы: числа в текстовом "
            "виде с запятой как десятичным разделителем, символ «?» вместо пропусков, "
            "смешение типов в колонках метрик."
        ),
        42: (
            "Очистка выполнена цепочкой встроенных компонентов Loginom. Узел "
            "«5. Очистка полей (изменение)» приводит структуру полей после объединения "
            "к единообразным именам и типам. «6. Пропуски (TBGDataRecoveryEngine)» "
            "заполняет пропущенные значения. «7. Выбросы (TBGElimOutlierEngine)» "
            "ограничивает экстремальные значения по числовым метрикам. «9b. Числовые "
            "метрики» переводит строковые метрики в числа. «9c. Квантование ключевых "
            "метрик» строит пять порядковых бинов (0–4) для popularity, сборов и рейтинга "
            "фильма — поля q_popularity, q_streams, q_track_score. Контрольный снимок "
            "после квантования — qc_after_quantization.txt."
        ),
        43: (
            "Целевая переменная IsHit задаётся правилом: 1, если Worldwide Gross не ниже "
            "ста миллионов, иначе 0. Дополнительно применён отбор «9. Отбор TMDB "
            "popularity ≥ 30», чтобы оставить в анализе фильмы с заметной видимостью "
            f"и снизить шум на «хвосте» низкой популярности. В финальной рабочей выборке "
            f"около {N_TOTAL} строк (train {N_TRAIN}, test {N_TEST}, valid {N_VALID})."
        ),
        47: (
            "Кратко по блокам на рисунке 3. «qc_before_cleaning.txt» — экспорт до recovery. "
            "«5. Очистка полей» — приведение схемы. «6. Пропуски» — заполнение null. "
            "«7. Выбросы» — обработка выбросов. «8. IsHit» и «9. Отбор TMDB popularity ≥ 30» "
            "— целевая метка и фильтр. «9b» и «9c» — парсинг и квантование. "
            "«qc_after_quantization.txt» — снимок после подготовки признаков."
        ),
        50: (
            "Данные разделены на обучающую, тестовую и валидационную части. Узел "
            "«8. Разбиение 80/20 (stratify IsHit)» делает стратифицированное разбиение "
            "по целевому классу. Метки SAMPLE: «9. SAMPLE = train» и «10. SAMPLE = test». "
            "«11. Объединение train + test» собирает поток обратно. «11b. SAMPLE = valid» "
            f"выделяет из train около 18,75 % строк в валидацию ({N_VALID} фильмов). "
            "Флаг IsTestSet истинен для test и valid. Итоговые объёмы зафиксированы в "
            "файле data/данные.txt. Публичный узел «Movies dataset (публичный)» передаёт "
            "подготовленный набор в сценарий моделирования."
        ),
        54: (
            "Кратко по блокам на рисунке 4. «8. Разбиение 80/20» — stratified split. "
            "«9» и «10» — метки train/test. «11» и «11b» — сборка потока и valid. "
            "«данные.txt» — экспорт подготовленной выборки. «Movies dataset (публичный)» "
            "— выход в моделирование."
        ),
        57: (
            "В сценарии «Построение модели» импортируется файл data/данные.txt, "
            "сформированный модулем подготовки. «Очистка числовых полей» — дополнительная "
            "страховка от нечисловых значений перед масштабированием. Блок «↶meta-scaling» "
            "разделяет поток по SAMPLE и обучает преобразование только на train. "
            "«↶preprocessing.Scaler» применяет StandardScaler к числовым признакам "
            "(id=OBJECT, target=IsHit). Узлы «NaN после scaling» заменяют пустые значения "
            "после масштабирования на ноль."
        ),
        58: (
            "Модель LOF настраивается в «↶neighbors.LOF Novelty»: Local Outlier Factor, "
            "режим novelty, n_neighbors=20, contamination=0,15. Обучение — узлом "
            "«↶model.fitter (выполнение)» на train; скоринг — на потоке test+valid."
        ),
        62: (
            "Кратко по блокам на рисунке 5. «Импорт data/данные.txt» — вход из Unit_0. "
            "«Очистка числовых полей» — финальный парсинг. «↶meta-scaling» + "
            "«↶preprocessing.Scaler» — нормализация только по train. «NaN после scaling» "
            "— замена null. «↶neighbors.LOF Novelty» + «↶model.fitter» — обучение и "
            "применение LOF."
        ),
        66: (
            "LOF выдаёт метку outlier_label: −1 — аномалия, +1 — норма. Для сравнения "
            "с IsHit в узле «Метки для metrics (0/1)» аномалия переводится в 1, норма в 0. "
            f"Узел «Только test для metrics» оставляет строки с IsTestSet=True (test и valid, "
            f"всего {N_TEST_VALID} записей). Компонент «↶classification metrics (выполнение)» "
            "считает precision, recall, F1, accuracy и матрицу ошибок."
        ),
        72: (
            f"По результатам на подвыборке test + valid (IsTestSet=True, {N_TEST_VALID} фильма) "
            "компонент classification metrics выводит precision, recall, F1 и матрицу ошибок "
            "(см. рисунки 6–7). Интерпретация: LOF в постановке «аномалия = блокбастер» "
            "ищет фильмы с необычными комбинациями кассовых сборов, популярности TMDB и "
            "социальных метрик. Статистическая аномалия не совпадает один в один с "
            "бизнес-определением IsHit по порогу сборов, поэтому метрики отражают "
            "компромисс между полнотой и точностью."
        ),
        73: (
            "Для учебной задачи важен полный цикл: ETL из Access, подготовка признаков, "
            "масштабирование, обучение LOF, расчёт метрик и сопоставление с целевой "
            "меткой блокбастера."
        ),
        77: (
            "Помимо LOF реализована ветка логистической регрессии sklearn: тот же блок "
            "meta-scaling и scaler, затем «↶linear_model.LogisticRegression (LR)» и "
            "«↶model.fitter (LR sklearn)». На test метрики LR выводятся в узле "
            "classification metrics (см. рисунок 8)."
        ),
        78: (
            f"Скоринг на новых данных выполнен отдельной веткой: «Импорт score CSV» — "
            f"{N_SCORE} фильмов из prepared_movies_score.csv; «Подготовка score» — "
            "приведение полей к схеме train, SAMPLE=score; «LR predict score (Python)» — "
            "q-бины и prob0/prob1; «Отчёт scoring LR» — сводка в report_scoring_lr.txt "
            f"({N_SCORE} строк, среднее prob1 ≈ {PROB1_MEAN:.4f}, при пороге 0,5 "
            f"предсказано {N_HITS_SCORE_PRED} блокбастеров)."
        ),
        89: (
            "В подготовке данных настроены пользовательские отчёты по целевой метке IsHit "
            "и контрольные txt-снимки qc_before_cleaning.txt и qc_after_quantization.txt."
        ),
        83: "Рисунок 9 — Скоринг новых фильмов (prob1)",
    }
    for idx, text in blocks.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], text)

    # Fix duplicate heading 3.6 -> 3.6 Оценка качества модели LOF
    if len(doc.paragraphs) > 65:
        set_paragraph_text(doc.paragraphs[65], "3.6 Оценка качества модели LOF")


def update_title_tables(doc: Document) -> None:
    """Student / teacher title blocks — placeholders for manual fill."""
    if len(doc.tables) >= 3:
        t = doc.tables[2]
        if len(t.rows) >= 1 and len(t.rows[0].cells) >= 6:
            t.rows[0].cells[1].text = "____"
            t.rows[0].cells[5].text = "Ф. И. О."


def build() -> Path:
    shutil.copy2(TEMPLATE, OUT)
    doc = Document(OUT)
    replace_in_doc(doc)
    rewrite_body(doc)
    update_title_tables(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Report written: {path}")
    # quick QA dump
    doc = Document(path)
    banned = ("Spotify", "spotify", "трек", "артист", "прослушиван")
    issues = []
    for i, p in enumerate(doc.paragraphs):
        for b in banned:
            if b in p.text:
                issues.append(f"p{i}: {p.text[:120]}")
    if issues:
        print("WARN banned terms:")
        for x in issues:
            print(" ", x)
    else:
        print("QA: no banned music terms in paragraphs")
