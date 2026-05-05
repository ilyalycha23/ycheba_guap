# -*- coding: utf-8 -*-
"""Генерация отчётов ЛР4–ЛР6 (Word) для варианта 12 по скриптам lab4–lab6."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BASE = Path(__file__).resolve().parent
VARIANT = 12


def set_normal_font(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)


def add_title(doc: Document, n: int, topic: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Лабораторная работа №{n}")
    r.bold = True
    r.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(topic)
    r.font.size = Pt(14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Студент: _______________________\nГруппа: __________   Вариант: {VARIANT}")


def add_h(doc: Document, text: str):
    doc.add_heading(text, level=1)


def lab4():
    doc = Document()
    set_normal_font(doc)
    add_title(
        doc,
        4,
        "Запросы на языке SQL: агрегатные функции "
        "(глава 4 методического указания, п. 4.3, приложение)",
    )

    add_h(doc, "1. Цель работы")
    doc.add_paragraph(
        "Освоить использование агрегатных функций и группировки (GROUP BY, HAVING), "
        "а также операций над результатами SELECT в духе п. 4.2 методички "
        "(UNION, INTERSECT, EXCEPT и варианты с ALL для мультимножеств в PostgreSQL 15+)."
    )

    add_h(doc, "2. Исходные данные и подготовка")
    doc.add_paragraph(
        "Схему базы данных создаём скриптом lab1/01_schema.sql; тестовые строки загружаем "
        "скриптом lab2/01_insert_test_data.sql. СУБД — PostgreSQL 11+. "
        "Клиент — psql, pgAdmin или аналог. Перед выполнением запросов п. 4.3 желательно "
        "убедиться, что в базе есть приёмы, штатное расписание и дети (это даёт загрузка lab2). "
        "При необходимости расширить данные только для локальных экспериментов в каталоге lab4 "
        "можно использовать 04_extra_data_for_lr4.sql (он не трогает lab2)."
    )

    add_h(doc, "3. Теоретическая часть (кратко по методичке)")
    doc.add_paragraph(
        "П. 4.1 — агрегаты AVG, SUM, COUNT, MIN, MAX; группировка по столбцам; отбор групп через HAVING. "
        "П. 4.2 — теоретико-множественные операции над результатами запросов: объединение, пересечение, "
        "разность; различие между операциями над множеством (дубликаты схлопываются) и над мультимножеством "
        "(кратности учитываются в UNION ALL, INTERSECT ALL, EXCEPT ALL там, где СУБД поддерживает ALL)."
    )

    add_h(doc, "4. Ход работы: запросы г)—е) варианта 12 (файл 01_queries_gde.sql)")
    doc.add_paragraph(
        "Согласно приложению к методичке для варианта 12 реализованы три пункта."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("г) Средняя численность подразделений. ").bold = True
    p.add_run(
        "Под «численностью подразделения» принимаем сумму поля количество_единиц по всем строкам "
        "таблицы штатное_расписания для данного ид_подразделения. Внутренний запрос группирует штат "
        "по подразделениям и считает SUM; внешний — усредняет эти суммы функцией AVG и округляет "
        "результат (ROUND) для читаемой таблицы результатов."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("д) Инженеры с более чем пятью детьми. ").bold = True
    p.add_run(
        "Строим связь сотрудник—дети, отбираем сотрудников, для которых существует приём "
        "на должность с категорией «инженер» (EXISTS по прием и должность). "
        "GROUP BY по сотруднику, HAVING COUNT(детей) > 5 сортируем по ФИО."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("е) Подразделения, где техников больше, чем инженеров. ").bold = True
    p.add_run(
        "Берём приёмы по подразделениям, к каждому присоединяем должность. "
        "В одной группе по подразделению считаем число DISTINCT сотрудников с категорией «техник» "
        "и отдельно — с категорией «инженер» (CASE внутри COUNT(DISTINCT …)). "
        "Условие на группы задаётся в HAVING."
    )
    doc.add_paragraph(
        "После выполнения каждого SELECT фиксируем скриншот или текстовый вывод (таблица результатов) "
        "и вставляем в отчёт."
    )

    add_h(doc, "5. Дополнительные агрегатные запросы (02_extra_aggregates.sql)")
    doc.add_paragraph(
        "В основных пунктах г)—е) для этого варианта MIN и MAX не использовались, поэтому "
        "дополнительно выполняем два запроса с минимумом и максимумом."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "По каждому подразделению: MIN и MAX даты приёма среди строк таблицы прием."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "По каждой должности из справочника: минимальное и максимальное количество_единиц "
        "в одной строке штатное_расписание для этой должности."
    )

    add_h(doc, "6. Операции п. 4.2 и мультимножество (03_set_operations_4_2.sql)")
    doc.add_paragraph(
        "Выполняем по очереди запросы из файла и сравниваем число строк."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "UNION двух выборок ФИО (например, на «И%» и на «%ов%») и UNION ALL на тех же данных: "
        "для совпадающих фамилий UNION оставляет одну строку, UNION ALL — две."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "INTERSECT: идентификаторы сотрудников, которые числятся в подразделениях с ид 2 и 4 "
        "(по таблице прием)."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "EXCEPT: множество сотрудников из приёмов минус множество уволенных из увольнение."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "На PostgreSQL 15+ — примеры INTERSECT ALL и EXCEPT ALL на конструкции VALUES для иллюстрации правил "
        "кратностей из п. 4.2 методички."
    )

    add_h(doc, "7. Выводы")
    doc.add_paragraph(
        "Выполнены агрегатные запросы с группировкой и HAVING для варианта 12; отработаны MIN/MAX "
        "на связанных с приёмами и штатом данных; продемонстрированы UNION/UNION ALL, INTERSECT, EXCEPT "
        "и при необходимости варианты ALL."
    )

    add_h(doc, "8. Приложения")
    doc.add_paragraph(
        "Листинги: lab4/01_queries_gde.sql, 02_extra_aggregates.sql, 03_set_operations_4_2.sql; "
        "при использовании — 04_extra_data_for_lr4.sql; схема и данные: lab1/01_schema.sql, "
        "lab2/01_insert_test_data.sql. Таблицы результатов выполнения запросов (скриншоты или вывод)."
    )

    doc.save(BASE / "lab4" / "Отчет_ЛР4_вариант12.docx")


def lab5():
    doc = Document()
    set_normal_font(doc)
    add_title(
        doc,
        5,
        "Подзапросы, [NOT] EXISTS и DML (глава 5 методического указания, п. 5.4)",
    )

    add_h(doc, "1. Цель работы")
    doc.add_paragraph(
        "Научиться формулировать условия через EXISTS и NOT EXISTS и вложенные подзапросы; "
        "использовать подзапросы в командах INSERT, UPDATE и DELETE; сопоставить логику п. 4.2 "
        "с формулировками через EXISTS; закрепить правило проверки NULL (IS NULL вместо = NULL)."
    )

    add_h(doc, "2. Подготовка базы")
    doc.add_paragraph(
        "Выполнить последовательно: lab1/01_schema.sql, lab2/01_insert_test_data.sql. "
        "При необходимости для более «насыщенной» базы перед ЛР5 — lab4/04_extra_data_for_lr4.sql "
        "(это необязательно для всех запросов)."
    )

    add_h(doc, "3. Ход работы: запросы ж)—и) варианта 12 (01_queries_zhi.sql)")
    doc.add_paragraph(
        "Перед запросом по пункту «з)» обязательно выполнить lab5/04_extra_data_lr5.sql: он добавляет "
        "строки в штатное_расписание для подразделения ОГК (ид 3), чтобы в данных появилось подразделение "
        "с полным набором должностей с ид 1–4. Скрипт идемпотентен (предварительно удаляет строки с ид 7, 8)."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("ж) Подразделения без совместителей. ").bold = True
    p.add_run(
        "Выводим все подразделение, для которых NOT EXISTS ни одной строки прием с совместительство = B'1'. "
        "Так корректно выражается «нигде не встречается» через вложенный подзапрос."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("з) Все должности представлены в штатном расписании подразделения. ").bold = True
    p.add_run(
        "Логика: для каждого подразделения не должно существовать такой должности из набора {1,2,3,4}, "
        "чтобы для неё NOT EXISTS строки штатное_расписание для этого подразделения. "
        "Двойное NOT EXISTS задаёт универсальную кванторную форму («для каждой должности найдётся штатная строка»)."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("и) У сотрудника все дети одного пола. ").bold = True
    p.add_run(
        "Требование «есть дети» — EXISTS по дети. Одновременно исключаем пары разных полов среди детей "
        "этого же сотрудника: NOT EXISTS по декартову произведению двух строк дети с условием ид1 < ид2 "
        "и пол1 <> пол2. Результат сортируется по ФИО."
    )
    doc.add_paragraph(
        "Результат каждого запроса фиксируем в приложении отчёта (скриншот или копирование вывода из psql)."
    )

    add_h(doc, "4. DML с подзапросами (02_01_insert, 02_02_update, 02_03_delete)")
    doc.add_paragraph(
        "Отдельные файлы с временными таблицами (TEMP), чтобы не портить учебную схему; в pgAdmin удобно по одному F5 на файл."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "CREATE TEMP lr5_tmp_ins; INSERT выбирает ид всех сотрудников, у которых есть дети "
        "(EXISTS из подзапроса к дети)."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "lr5_tmp_upd: строки создаются по всем подразделениям, затем UPDATE ставит счётчик равным "
        "подзапросу COUNT(*) из прием по соответствующему ид_подразделения."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "lr5_tmp_del: наполняется ид сотрудников с ид<=3; DELETE удаляет из этой временной таблицы тех, "
        "для кого EXISTS строка увольнение (имитация удаления только «ещё работающих» из набора)."
    )
    doc.add_paragraph(
        "В каждом файле в конце уже есть свой SELECT — сохраняем три скрина/вывода в отчёт (см. также 02_dml_subqueries.sql)."
    )

    add_h(doc, "5. EXISTS и аналоги INTERSECT/EXCEPT, про NULL (03_exists_vs_set_nulls.sql)")
    doc.add_paragraph(
        "Первый запрос — те же ид приёмы в подразделениях 2 и 4, что и INTERSECT в ЛР4, но через EXISTS: "
        "для каждого сотрудника из приёмов деп. 2 проверяем наличие приёма в деп. 4."
    )
    doc.add_paragraph(
        "Второй — аналог EXCEPT: работники из прием, для которых NOT EXISTS увольнение с тем же "
        "ид_сотрудника."
    )
    doc.add_paragraph(
        "Затем два запроса к VALUES (NULL, 1): с условием x IS NULL даёт строку; с условием x = NULL "
        "в стандартной семантике SQL даёт пустой результат (сравнение с NULL не TRUE). Комментируем вывод для отчёта."
    )

    add_h(doc, "6. Выводы")
    doc.add_paragraph(
        "EXISTS позволяет выразить «есть хотя бы один» и через отрицание — «нет ни одного», что удобно "
        "для сложных кванторных формулировок. Подзапросы в DML локализуют правила отбора строк в самой команде "
        "модификации. Эквивалент INTERSECT/EXCEPT через EXISTS не завязан на операции множества и переносимее "
        "в СУБД без полного набора ключевых слов ALL. Корректная проверка неизвестного значения выполняется "
        "только через IS NULL / IS NOT NULL."
    )

    add_h(doc, "7. Приложения")
    doc.add_paragraph(
        "Файлы: lab5/01_queries_zhi.sql, 02_01_insert.sql, 02_02_update.sql, 02_03_delete.sql (указатель: 02_dml_subqueries.sql), 03_exists_vs_set_nulls.sql, "
        "04_extra_data_lr5.sql; lab1/01_schema.sql, lab2/01_insert_test_data.sql; при необходимости "
        "lab4/04_extra_data_for_lr4.sql. Результаты выполнения и скриншоты."
    )

    doc.save(BASE / "lab5" / "Отчет_ЛР5_вариант12.docx")


def lab6():
    doc = Document()
    set_normal_font(doc)
    add_title(
        doc,
        6,
        "Хранимые процедуры и функции (глава 6 методического указания, п. 6.4)",
    )

    add_h(doc, "1. Цель работы")
    doc.add_paragraph(
        "Освоить создание процедур (CREATE PROCEDURE … PL/pgSQL) и функций (CREATE FUNCTION) в PostgreSQL: "
        "вставка с проверкой и автодобавлением справочника, ручной каскад удаления, безопасное удаление "
        "структуры подразделения, накопление агрегатов во временную таблицу, управляющие конструкции "
        "(IF, FOR, WHILE), скалярная и табличная функции."
    )

    add_h(doc, "2. Подготовка")
    doc.add_paragraph(
        "Выполнить lab1/01_schema.sql и lab2/01_insert_test_data.sql. При повторной отладке объектов ЛР6 "
        "можно использовать lab6/09_drop_lr6_objects.sql (по указанию преподавателя). PostgreSQL версии "
        "11+ нужен для поддержки CALL и процедур."
    )

    add_h(doc, "3. Ход работы по файлам lab6")
    doc.add_paragraph(
        "Ниже — порядок создания объектов и смысл каждого. Выполняем скрипты 01–08 в том же порядке, "
        "что указан в readme папки lab6; затем демонстрации."
    )

    doc.add_paragraph(style="List Bullet").add_run(
        "01_proc_insert_with_directory.sql — процедура lr6_insert_hire_with_position: проверка существования "
        "сотрудника и подразделения; при отсутствии строки должности заданного ид — автоматическая вставка "
        "в справочник должность с запасными наименованием и категорией; новый ид приёма через MAX+1; "
        "INSERT в таблицу прием."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "02_proc_delete_employee_cascade.sql — lr6_delete_employee: вручную удаляются зависимости "
        "(перевод, увольнение, прием, дети), затем сам сотрудник; имитация каскада при ограничениях NO ACTION "
        "(осторожно с реальной БД — либо тестовая копия, либо ROLLBACK в транзакции)."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "03_proc_delete_subdivision_safe.sql — lr6_delete_leaf_subdept: удаление возможно только для "
        "«листового» подразделения без дочерних записей в подразделение, без штатного расписания и без "
        "приёмов иначе RAISE EXCEPTION."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "04_fn_aggregate.sql — функция lr6_sum_staff_units(ид_подразделения) возвращает сумму количества_единиц "
        "из штатное_расписание (SQL-модификатор STABLE)."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "05_proc_stats_temp.sql — lr6_fill_stats_temp пересоздаёт TEMP TABLE lr6_stat и заполняет по каждому "
        "подразделению: число строк штата, сумму единиц (через lr6_sum_staff_units), число различных "
        "работников по приёмам."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "06_batch_control_flow.sql — анонимный блок DO с FOR по подразделениям, счётчиком подразделений "
        "с приёмами (IF EXISTS), счётчиком выводится через RAISE NOTICE; далее цикл WHILE для демонстрации."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "07_fn_scalar.sql — lr6_child_count(ид_сотрудника) возвращает число записей в дети."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "08_fn_table.sql — lr6_direct_subdepts(родитель): табличная функция прямых дочерних подразделений; "
        "при NULL в аргументе — выборка корневых (ид_родительское IS NULL)."
    )

    add_h(doc, "4. Демонстрация выполнения")
    doc.add_paragraph(
        "Зафиксировать в отчёте вывод psql или скриншоты."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "CALL lr6_insert_hire_with_position(1, 2, 90, CURRENT_DATE, B'0') — при отсутствии должности 90 "
        "появится строка в должность и новая строка прием; проверить SELECT."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "Удаление сотрудника/подразделения — только по согласованию; предпочтительно 10_demo_transaction.sql "
        "с BEGIN … ROLLBACK, чтобы не ломать общие данные."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "SELECT lr6_sum_staff_units(2); CALL lr6_fill_stats_temp(); SELECT * FROM lr6_stat ORDER BY 1;"
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "Выполнить \\i lab6/06_batch_control_flow.sql и сохранить NOTICE."
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "SELECT с.ид, lr6_child_count(с.ид) FROM сотрудник с ORDER BY 1;"
    )
    doc.add_paragraph(style="List Bullet").add_run(
        "SELECT * FROM lr6_direct_subdepts(1); при необходимости проверить вызов с NULL."
    )

    add_h(doc, "5. Выводы")
    doc.add_paragraph(
        "Процедуры инкапсулируют проверку целостности и повторяемый сценарий DML, что упрощает клиентские "
        "приложения и снижает риск ошибок. Временная таблица lr6_stat живёт в сессии и не конфликтует с другими "
        "пользователями. Табличная функция удобна, когда результат — набор строк с той же структурой для JOIN "
        "или повторного использования в запросах."
    )

    add_h(doc, "6. Приложения")
    doc.add_paragraph(
        "lab6/01–08,10,09 как в readme; при необходимости lab1/01_schema.sql, lab2/01_insert_test_data.sql "
        "и тексты объектов как листинги в конце работы."
    )

    doc.save(BASE / "lab6" / "Отчет_ЛР6_вариант12.docx")


if __name__ == "__main__":
    lab4()
    lab5()
    lab6()
    for sub in ("lab4", "lab5", "lab6"):
        print("OK:", BASE / sub / f"Отчет_ЛР{sub[-1]}_вариант{VARIANT}.docx")
